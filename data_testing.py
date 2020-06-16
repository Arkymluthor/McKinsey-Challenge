import glob
import os

# Get the date of execution
import datetime
date_generated = datetime.datetime.now()

from platform import python_version
# use python_version() to get the version.

import pandas as pd

# set maximum number of columns to display in notebook
pd.set_option('display.max_columns', 999)

# To check whether a column is numeric type
from pandas.api.types import is_numeric_dtype

# To check whether a column is object/string type
from pandas.api.types import is_string_dtype

import numpy as np
from matplotlib.patches import Rectangle

# Import the graph packages
import matplotlib.pyplot as plt
plt.rcParams.update({'figure.max_open_warning': 0})
%matplotlib inline
import seaborn as sns

# This library is required to generate the MS Word document
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH  #used to align str(number) in cells

for infile in glob.glob(r"C:\Users\malik\Udacity Bertelsmann Scholarship\Intro to Neural Networks\Incident Data\*.xlsx"):
    name = os.path.basename(infile)
    filename, file_extension = os.path.splitext(name)
    data = pd.read_excel(infile)
    data = data.dropna(how='all', axis=1)

    # Number of rows of the DPD will be the count of columns in the raw date dataframe
    # Since it there will be one row for each column
    no_of_rows = len(data.columns)


    # Constructing the data_qlt_df dataframe and pre-assigning and columns
    # Pre-assigning the number of rows the dataframe would have is memory and processing efficient
    # This is a better approach than continuous append or concat operation to dataframe

    data_qlt_df = pd.DataFrame(index=np.arange(0, no_of_rows), \
                                columns=('column_name', 'col_data_type'
                                         ,'non_null_values', \
                                         'unique_values_count', 'column_dtype')
                              )


    # Add rows to the data_qlt_df dataframe
    for ind, cols in enumerate(data.columns):
        # Count of unique values in the column
        col_unique_count = data[cols].nunique()

        data_qlt_df.loc[ind] = [cols, \
                                data[cols].dtype, \
                                data[cols].count(), \
                                col_unique_count, \
                                cols + '~'+ str(data[cols].dtype)
                                ]

    # Use describe() to get column stats of raw dataframe
    # This will be merged with the DPD
    raw_num_df = data.describe().T.round(2)

    #----- Key Step ---------------
    # Merging the df.describe() output with rest of the info to create a single Data Profile Dataframe
    data_qlt_df = pd.merge(data_qlt_df, raw_num_df, how='left', left_on='column_name', right_index=True)


    # Calculate percentage of non-null values over total number of values
    data_qlt_df['%_of_non_nulls'] = (data_qlt_df['non_null_values']/data.shape[0])*100

    # Calculate null values for the column
    data_qlt_df['null_values'] = data.shape[0] - data_qlt_df['non_null_values']

    # Calculate percentage of null values over total number of values
    data_qlt_df['%_of_nulls'] = 100 - data_qlt_df['%_of_non_nulls']

    # Calculate the count of each data type
    data_qlt_df["dtype_count"] = data_qlt_df.groupby('col_data_type')["col_data_type"].transform('count')

    # Calculate the total count of column values
    data_qlt_df["count"] = data_qlt_df['null_values'] + data_qlt_df['non_null_values']

    # Reorder the Data Profile Dataframe columns
    data_qlt_df = data_qlt_df[
                                ['column_name', 'col_data_type',\
                                 'dtype_count', 'non_null_values', '%_of_non_nulls',\
                                 'null_values', '%_of_nulls', 'unique_values_count', 'count', 'mean', 'std', 'min', '25%',\
                                 '50%', '75%', 'max']
                             ]

    # Get the list of numeric columns from raw dataframe
    # need this: from pandas.api.types import is_numeric_dtype
    # get numeric columns which are not empty
    num_cols = [cols for cols in data.columns if is_numeric_dtype(data[cols]) and len(data[cols].dropna())>0]

    iter_len = len(num_cols)

    # For each numeric column in the list
    for x, col_name in enumerate(num_cols):
        #print(x+1, " of ", iter_len, " completed   ",  col_name)

        # Create a copy of the column values without nulls or NA
        no_null_col = data[col_name].dropna()


        # Calculate the 95 percentile of the values
        q25 = np.percentile(no_null_col, 25)
        q75 = np.percentile(no_null_col, 75)
        q95 = np.percentile(no_null_col, 95)

        # Plot the graphs
        fig3 = plt.figure(figsize=(20,15))
        fig3.suptitle("Profile of column  " + col_name, fontsize=25)  #Title for the whole figure
        plt.subplots_adjust(wspace=0.4, hspace=0.35)

        ax1 = fig3.add_subplot(2,3,1)
        ax1.set_title("Box plot for all the values", fontsize=20)
        plt.setp(ax1.get_xticklabels(), ha="right", rotation=35)
        plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)
        ax1.boxplot(no_null_col)

        ax1 = fig3.add_subplot(2,3,2)
        ax1.set_title("Distribution of all values", fontsize=20)
        plt.setp(ax1.get_xticklabels(), ha="right", rotation=35, fontsize=15)
        plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)
        ax1.hist(no_null_col)

        ax1 = fig3.add_subplot(2,3,3)
        ax1.set_title("Boxplot for quartiles (all values)", fontsize=20)
        if len(no_null_col.value_counts()) >= 4:
            data[u'quartiles'] = pd.qcut(
                            data[col_name],
                            4, duplicates='drop')
            data.boxplot(column= col_name, by=u'quartiles', ax = ax1)
        plt.setp(ax1.get_xticklabels(), ha="right", rotation=35, fontsize=15)
        plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)

        ax1 = fig3.add_subplot(2,3,4)
        ax1.set_title("Box plot without outliers", fontsize=20)
        plt.setp(ax1.get_xticklabels(), ha="right", rotation=35, fontsize=15)
        plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)
        ax1.boxplot(no_null_col, showfliers=False)

        ax1 = fig3.add_subplot(2,3,5)
        ax1.set_title("Violin plot (<95% percentile)", fontsize=20)
        plt.setp(ax1.get_xticklabels(), ha="right", rotation=35, fontsize=15)
        plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)
        ax1.violinplot(no_null_col[no_null_col <= q95])


        #Histogram with bin ranges, counts and percentile color
        ax1 = fig3.add_subplot(2,3,6)
        ax1.set_title("Histogram (<95% percentile)", fontsize=20)
        plt.setp(ax1.get_xticklabels(), ha="right", rotation=35, fontsize=15)
        plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)

        # Take only the data less than 95 percentile
        datas = no_null_col[no_null_col <= q95]

        # Colours for different percentiles
        perc_25_colour = 'gold'
        perc_50_colour = 'mediumaquamarine'
        perc_75_colour = 'deepskyblue'
        perc_95_colour = 'peachpuff'

        '''
        counts  = numpy.ndarray of count of data ponts for each bin/column in the histogram
        bins    = numpy.ndarray of bin edge/range values
        patches = a list of Patch objects.
                each Patch object contains a Rectnagle object.
                e.g. Rectangle(xy=(-2.51953, 0), width=0.501013, height=3, angle=0)
        '''
        counts, bins, patches = ax1.hist(datas, bins=10, facecolor=perc_50_colour, edgecolor='gray')

        # Set the ticks to be at the edges of the bins.
        ax1.set_xticks(bins.round(2))
        plt.xticks(rotation=70, fontsize=15)

        # Change the colors of bars at the edges
        for patch, leftside, rightside in zip(patches, bins[:-1], bins[1:]):
            if rightside < q25:
                patch.set_facecolor(perc_25_colour)
            elif leftside > q95:
                patch.set_facecolor(perc_95_colour)
            elif leftside > q75:
                patch.set_facecolor(perc_75_colour)

        # Calculate bar centre to display the count of data points and %
        bin_x_centers = 0.5 * np.diff(bins) + bins[:-1]
        bin_y_centers = ax1.get_yticks()[1] * 0.25

        # Display the the count of data points and % for each bar in histogram
        for i in range(len(bins)-1):
            bin_label = "{0:,}".format(counts[i]) + "  ({0:,.2f}%)".format((counts[i]/counts.sum())*100)
            plt.text(bin_x_centers[i], bin_y_centers, bin_label, rotation=90, rotation_mode='anchor')

        #create legend
        handles = [Rectangle((0,0),1,1,color=c,ec="k") for c in [perc_25_colour, perc_50_colour, perc_75_colour, perc_95_colour]]
        labels= ["0-25 Percentile","25-50 Percentile", "50-75 Percentile", ">95 Percentile"]
        plt.legend(handles, labels, bbox_to_anchor=(0.5, 0., 0.85, 0.99))


        fig3.suptitle("Profile of column  " + col_name, fontsize=25)  #Title for the whole figure
        fig_name = 'fig_' + col_name
        fig3.savefig(fig_name, dpi=50)
        plt.close('all')

    #     plt.show()

    data.drop(u'quartiles', axis=1, inplace=True)

    # Get the list of object columns from raw dataframe
    # get object columns which are not empty
    #obj_cols = [cols for cols in data.columns if is_string_dtype(data[cols]) and len(data[cols].dropna())>0]
    obj_cols = [cols for cols in data.columns if not is_numeric_dtype(data[cols]) and len(data[cols].dropna())>0]

    iter_len = len(obj_cols)


    # For each object column in the list
    for x, col_name in enumerate(obj_cols):
        #print(x+1, " of ", iter_len, " completed   ",  col_name)

        # Create a copy of the column values without nulls or NA
        no_null_col = data[col_name].dropna()

        values_freq_threshold = 25
        col_unique_count = data[col_name].nunique()

        # If unique values count is below the threshold value then store the details of unique values
        col_unique_vals = data[col_name].value_counts(normalize=True, sort=True)

        # Plot the graphs
        fig4 = plt.figure(figsize=(20,7))
        fig4.suptitle("Profile of column  " + col_name, fontsize=25)  #Title for the whole figure
        plt.subplots_adjust(wspace=0.4, hspace=0.35, bottom=0.35)

        ax1 = fig4.add_subplot(1,1,1)
        ax1.set_title("Bar chart for top 25 values", fontsize=20)
        plt.setp(ax1.get_xticklabels(), ha="right", rotation=45, fontsize=15)
        plt.setp(ax1.get_yticklabels(), ha="right", fontsize=15)

        col_unique_vals.head(values_freq_threshold).sort_values(ascending=False).plot.bar()
        plt.xticks(rotation=75)
        for p in ax1.patches:
            ax1.annotate(str(round(p.get_height(),2)), (p.get_x() * 1.005, p.get_height() * 1.005), fontsize=15)

        fig4.suptitle("Profile of column  " + col_name, fontsize=25)  #Title for the whole figure
        fig_name = 'fig_' + col_name
        fig4.savefig(fig_name, dpi= 50)

        plt.close('all')
    #     plt.show()

    ## Plot correlation
    f, ax = plt.subplots(figsize=(15, 10))
    plt.subplots_adjust(bottom=0.35)
    plt.autoscale()

    corr_data = data.corr()
    sns.heatmap(corr_data,
                mask=np.zeros_like(corr_data, dtype=np.bool),
                cmap=sns.diverging_palette(20, 220, as_cmap=True),
                vmin=-1, vmax=1,
                square=True,
                ax=ax)

    fig_name = 'fig_cor_plot.png'
    f.savefig(fig_name,  dpi=70)
    # plt.show()
    plt.close('all')

    # Make sure you have the docx package and it is imported
    # see the environment setup section

    #Create Document object
    document = Document()

    # Add Title
    document.add_heading('Data Profile -', 0)
    document.add_heading(filename, 0)

    # COver page paragraph
    p = document.add_paragraph('The main objective of this document is ')
    #p.add_run('only').bold = True
    p.add_run(' to understand raw data profile. i.e. data type, min & max values, ranges, unique values, etc.')
    p = document.add_paragraph('This present document only shows all the causes of failure without the details and what measures \
    taking.')
    #p = document.add_paragraph('')
    #p.add_run('The code is largely kept generic so that it could be used with any shape of data.').italic = True

    document.add_page_break()
    document.add_heading('Columns Data Profile Summary', 0)

    # Page 4
    p = document.add_paragraph(' ')

    # Heading 1
    document.add_heading('Dataset shape', level=1)

    table = document.add_table(rows=2, cols=2, style = 'Medium Shading 1 Accent 3')

    # Header row
    cell = table.cell(0, 0)
    cell.text = 'No.of rows'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True

    cell = table.cell(0, 1)
    cell.text = 'No.of columns'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True

    # Values
    cell = table.cell(1, 0)
    cell.text = F'{data.shape[0] :,}'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = False

    cell = table.cell(1, 1)
    cell.text = F'{data.shape[1] :,}'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = False

    # Page 4a
    # document.add_page_break()
    p = document.add_paragraph(' ')

    # Heading 1
    document.add_heading('Dataframe columns summary', level=1)

    # Rehsape the column data type dataframe into form that can be printed in MS Word
    datas = round(data_qlt_df[['column_name','col_data_type', 'non_null_values', 'null_values', 'count']], 2)

    # add a table to the end and create a reference variable
    # extra row is so we can add the header row
    table = document.add_table(datas.shape[0]+1, datas.shape[1], style='Medium Shading 1 Accent 3')

    # add the header rows.
    for j in range(datas.shape[1]):

        #header row first two columns
        if j <= 1:
            cell = table.cell(0, j)
            cell.text = F'{datas.columns[j]}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = True
        else:
            cell = table.cell(0, j)
            cell.text = F'{datas.columns[j]}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = True
            cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT


    # add the rest of the data frame
    for i in range(datas.shape[0]):
        for j in range(datas.shape[1]):
            if j <= 1:
                cell = table.cell(i+1, j)
                cell.text = F'{datas.values[i,j]}'
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.size = Pt(11)
                cell_font.bold = False
            else:
                cell = table.cell(i+1, j)
                cell.text = F'{datas.values[i,j] :,}'
                cell_font = cell.paragraphs[0].runs[0].font
                cell_font.size = Pt(11)
                cell_font.bold = False
                cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT

    document.add_page_break()
    document.add_heading('Data correlation plot', 0)

    p = document.add_paragraph('')

    document.add_picture('fig_cor_plot.png', height=Inches(6), width=Inches(6))

    document.add_page_break()
    document.add_heading('Column Data Profile Details', 0)

    # ind = 1  # to be taken from iterrows loop later
    for ind in range(data_qlt_df.shape[0]):
        document.add_page_break()

        # Create table for column profile details
        table = document.add_table(rows=6, cols=6, style = 'Medium Shading 1 Accent 3' )

        # Merge cells in header row for COlumn Name
        for y in range(len(table.rows[0].cells)-1):
            a = table.cell(0,y)
            b = table.cell(0,y+1)
            a.merge(b)

        # Merge cells in detail rows spanning 2 cells x 3
        for row in range(1,6):
            a = table.cell(row,0)
            b = table.cell(row,1)
            a.merge(b)
            a = table.cell(row,2)
            b = table.cell(row,3)
            a.merge(b)
            a = table.cell(row,4)
            b = table.cell(row,5)
            a.merge(b)


        #*** ADD VALUES TO TABLE  ***#
        # Cell 0,0 (merged 6 cells): Header - Column Name
        cell = table.cell(0, 0)
        cell.text = data_qlt_df["column_name"][ind]
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(15)
        cell_font.bold = True

        # Cell 1,0: Blank
        cell = table.cell(1, 1)
        cell.text = "TBD Column :\n"
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        p = cell.paragraphs[0].add_run('no value')
        cell_font2 = cell.paragraphs[0].runs[1].font
        cell_font2.size = Pt(12)
        cell_font2.bold = False

        # Cell 1,0: Column data type
        cell = table.cell(1, 3)
        cell.text = 'Data Type : \n'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        p = cell.paragraphs[0].add_run(str(data_qlt_df["col_data_type"][ind]))
        cell_font2 = cell.paragraphs[0].runs[1].font
        cell_font2.size = Pt(12)
        cell_font2.bold = False

        # Cell 1,1: Count of toal values in the column
        cell = table.cell(1, 5)
        cell.text = 'Values Count : \n'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        p = cell.paragraphs[0].add_run(F'{data_qlt_df["count"][ind] :,.0f}')
        cell_font2 = cell.paragraphs[0].runs[1].font
        cell_font2.size = Pt(11)
        cell_font2.bold = False

        # Cell 2,0: Count of unique values in the column
        cell = table.cell(2, 1)
        cell.text = 'Unique Values Count : \n'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        unique_per = (data_qlt_df["unique_values_count"][ind] / data_qlt_df["count"][ind]) * 100
        p = cell.paragraphs[0].add_run(F'{data_qlt_df["unique_values_count"][ind] :,.0f}' + "   " + F'({unique_per :,.2f}%)' )
        cell_font2 = cell.paragraphs[0].runs[1].font
        cell_font2.size = Pt(11)
        cell_font2.bold = False

        # Cell 2,1: Count of non-null values in the column
        cell = table.cell(2, 3)
        cell.text = 'Non-Null Values Count : \n'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        p = cell.paragraphs[0].add_run(F'{data_qlt_df["non_null_values"][ind] :,.0f}' + "   " + F' ({data_qlt_df["%_of_non_nulls"][ind]  :,.2f}%)' )
        cell_font2 = cell.paragraphs[0].runs[1].font
        cell_font2.size = Pt(11)
        cell_font2.bold = False

        # Cell 2,2: Count of null values in the column
        cell = table.cell(2, 5)
        cell.text = 'Null Values Count : \n'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        p = cell.paragraphs[0].add_run(F'{data_qlt_df["null_values"][ind]  :,.0f}' + "   " + F' ({data_qlt_df["%_of_nulls"][ind]  :,.2f}%)' )
        cell_font2 = cell.paragraphs[0].runs[1].font
        cell_font2.size = Pt(11)
        cell_font2.bold = False

        # Cell 3,0: Min of values in the column
        cell = table.cell(3, 1)
        cell.text = 'Min : \n'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        p = cell.paragraphs[0].add_run(F'{data_qlt_df["min"][ind]  :,.2f}' )
        cell_font2 = cell.paragraphs[0].runs[1].font
        cell_font2.size = Pt(11)
        cell_font2.bold = False

        # Cell 3,1: Mean of values in the column
        cell = table.cell(3, 3)
        cell.text = 'Mean :  \n'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        p = cell.paragraphs[0].add_run(F'{data_qlt_df["mean"][ind] :,.2f}' )
        cell_font2 = cell.paragraphs[0].runs[1].font
        cell_font2.size = Pt(11)
        cell_font2.bold = False

        # Cell 3,3: Max of values in the column
        cell = table.cell(3, 5)
        cell.text = 'Max : \n'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        p = cell.paragraphs[0].add_run(F'{data_qlt_df["max"][ind]  :,.2f}' )
        cell_font2 = cell.paragraphs[0].runs[1].font
        cell_font2.size = Pt(11)
        cell_font2.bold = False

        # Cell 4,1: 25th Percentile of values in the column
        cell = table.cell(4, 1)
        cell.text = '25th Percentile : \n'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        p = cell.paragraphs[0].add_run(F'{data_qlt_df["25%"][ind]  :,.2f}' )
        cell_font2 = cell.paragraphs[0].runs[1].font
        cell_font2.size = Pt(11)
        cell_font2.bold = False

        # Cell 4,2: 50th Percentile of values in the column
        cell = table.cell(4, 3)
        cell.text = '50th Percentile : \n'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        p = cell.paragraphs[0].add_run(F'{data_qlt_df["50%"][ind]  :,.2f}' )
        cell_font2 = cell.paragraphs[0].runs[1].font
        cell_font2.size = Pt(11)
        cell_font2.bold = False

        # Cell 4,3: 75th Percentile of values in the column
        cell = table.cell(4, 5)
        cell.text = '75th Percentile : \n'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        p = cell.paragraphs[0].add_run(F'{data_qlt_df["75%"][ind]  :,.2f}' )
        cell_font2 = cell.paragraphs[0].runs[1].font
        cell_font2.size = Pt(11)
        cell_font2.bold = False


        p = document.add_paragraph(' ')
        p = document.add_paragraph(' ')

        fig_name = 'fig_' + data_qlt_df['column_name'][ind] + '.png'
        document.add_picture(fig_name, height=Inches(3.5), width=Inches(6))

    # save the doc
    document.save(r'C:\Users\malik\Udacity Bertelsmann Scholarship\Intro to Neural Networks\Incident Data\Report_%s.docx' %filename)
    print("Document generated!")

for file in os.listdir(r'C:\Users\malik\Udacity Bertelsmann Scholarship\Intro to Neural Networks'):
    if file.endswith('.png'):
        os.remove(file)
